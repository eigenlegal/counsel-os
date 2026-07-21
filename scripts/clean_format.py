#!/usr/bin/env python3
"""Reformat a simple .docx to professional standards using consistent styles.

Usage:
    python3 clean_format.py <input.docx> <output.docx> [--template <template.docx>]

This script rebuilds the output document from body paragraphs and tables. It is
appropriate for simple letters, memos, and freshly drafted contracts, not
negotiated documents or files that must preserve comments, tracked changes,
hyperlinks, fields, images, footnotes/endnotes, or section-specific formatting.

Applies uniform 11pt Times New Roman, justified body text, bold headings with
proper spacing, native Word multilevel numbering, and 1-inch margins. Straight
quotes are converted to curly (smart) quotes.
The default template is legal-template.docx in the same directory as this script.

Numbering always ends up native (Word w:numPr), never literal text, and works
in one of two modes:

  mirror  Two or more headings carry literal section numbers, so the document
          is presumed to number itself deliberately. Levels are taken from
          each number's own depth ("1.1" is a level-2 section regardless of
          its heading style), and the literal number is stripped in favour of
          the native one. Numbering normalises to a sequential run, so
          duplicate, backwards and skipped numbers are corrected and reported
          in results["corrections"]; an exhibit or schedule heading opens a
          fresh region that restarts at 1.

  auto    No literal numbering, so numbering is invented from heading styles
          (memos, letters).

In both modes a paragraph is only numbered if it is genuinely a section.
Draft legends, privilege banners, centred text and recital prose are styled
but never join the section counter.

Correcting numbers does not rewrite prose: if a heading is renumbered and the
body cites section numbers, a warning says so and the citations need checking
by hand.
"""

import json
import re
import sys
from pathlib import Path
from typing import NamedTuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

DEFAULT_FONT = "Times New Roman"
DEFAULT_SIZE = Pt(11)

# ---------------------------------------------------------------------------
# Numbering — multilevel 1 / 1.1 / 1.1.1
# ---------------------------------------------------------------------------

_NSMAP = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _make_level(ilvl: int, num_fmt: str, lvl_text: str,
                num_pos: int, tab_pos: int, wrap_pos: int = 0, start: int = 1):
    """Build one <w:lvl> element for an abstract numbering definition.

    num_pos:   where the number appears (twips from left margin)
    tab_pos:   tab stop after the number (where text begins on first line)
    wrap_pos:  where continuation lines wrap to (0 = left margin)

    The level carries its own <w:rPr>. Word renders the auto-number glyph
    using the numbering level's run properties, not the paragraph's, so
    without this the number falls back to docDefaults — the source template's
    theme font — and renders in Cambria or Calibri beside Times New Roman
    text.
    """
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), str(ilvl))

    st = OxmlElement("w:start")
    st.set(qn("w:val"), str(start))
    lvl.append(st)

    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    lvl.append(fmt)

    txt = OxmlElement("w:lvlText")
    txt.set(qn("w:val"), lvl_text)
    lvl.append(txt)

    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)

    pPr = OxmlElement("w:pPr")

    # Tab stop for clean gap after number
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(tab_pos))
    tabs.append(tab)
    pPr.append(tabs)

    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(wrap_pos))
    # Use firstLine to push the number right on the first line
    if num_pos > wrap_pos:
        ind.set(qn("w:firstLine"), str(num_pos - wrap_pos))
    pPr.append(ind)
    lvl.append(pPr)

    # w:rPr follows w:pPr in the w:lvl content model.
    lvl.append(_number_run_properties())

    return lvl


def _number_run_properties():
    """<w:rPr> pinning the auto-number glyph to the document font."""
    rPr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attr), DEFAULT_FONT)
    rPr.append(fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(DEFAULT_SIZE.pt * 2)))
    rPr.append(sz)
    return rPr


def _add_abstract_numbering(numbering_elm, abstract_num_id: int, levels):
    """Append one abstractNum definition to the numbering element."""
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_num_id))

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "multilevel")
    abstract_num.append(multi_level)

    for lvl in levels:
        abstract_num.append(lvl)

    numbering_elm.append(abstract_num)


def _add_num_instance(numbering_elm, num_id: int, abstract_num_id: int,
                      restart: bool = False, level_count: int = 4):
    """Append one w:num instance pointing at an abstract definition.

    Two w:num instances that share an abstractNum also share its counters
    unless the second carries explicit startOverride values, so `restart`
    is the only reliable way to make an exhibit begin again at 1.
    """
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract_ref)

    if restart:
        for ilvl in range(level_count):
            override = OxmlElement("w:lvlOverride")
            override.set(qn("w:ilvl"), str(ilvl))
            start_override = OxmlElement("w:startOverride")
            start_override.set(qn("w:val"), "1")
            override.append(start_override)
            num.append(override)

    numbering_elm.append(num)


def _add_numbering_definition(numbering_elm, abstract_num_id: int, num_id: int, levels):
    """Append one abstractNum + num pair to the numbering element."""
    _add_abstract_numbering(numbering_elm, abstract_num_id, levels)
    _add_num_instance(numbering_elm, num_id, abstract_num_id)


# Heading geometry, one (num_pos, tab_pos) pair per level, in twips.
# Numbers sit flush left with a modest gap to the text, which is the standard
# contract look; 1440 twips = 1 inch.
HEADING_GEOMETRY = [
    (0, 360),   # 1.        text at 0.25"
    (0, 720),   # 1.1       text at 0.50"
    (0, 1080),  # 1.1.1     text at 0.75"
    (0, 1440),  # 1.1.1.1   text at 1.00"
]

HEADING_LEVELS = len(HEADING_GEOMETRY)

# numId 1 is the body numbering region, numId 2 is lists, and each additional
# numbering region (one per exhibit/schedule) takes the next free id.
BODY_NUM_ID = 1
LIST_NUM_ID = 2
_FIRST_REGION_NUM_ID = 3


def create_numbering(doc, extra_regions: int = 0):
    """Create multilevel numbering definitions for headings and lists.

    Purges all pre-existing numbering definitions to avoid conflicts with
    built-in styles, then creates two fully independent definitions: one for
    headings (1, 1.1, 1.1.1) and one for bullet/numbered list items. Keeping
    them separate means list items never advance the heading counters (a
    3-item list under section 2 no longer pushes the next subsection to 2.4).

    `extra_regions` additional heading instances are created, each restarting
    at 1, for documents whose numbering restarts at an exhibit or schedule.

    Returns (heading_num_ids, list_num_id) where heading_num_ids[i] is the
    numId for numbering region i.
    """
    numbering_part = doc.part.numbering_part
    numbering_elm = numbering_part._element

    # Purge all existing abstractNum and num elements
    for child in list(numbering_elm):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("abstractNum", "num"):
            numbering_elm.remove(child)

    # --- Headings: 1. / 1.1. / 1.1.1. / 1.1.1.1 ---
    heading_levels = [
        _make_level(
            ilvl,
            "decimal",
            ".".join(f"%{n + 1}" for n in range(ilvl + 1)) + ".",
            num_pos=num_pos,
            tab_pos=tab_pos,
            wrap_pos=0,
        )
        for ilvl, (num_pos, tab_pos) in enumerate(HEADING_GEOMETRY)
    ]
    _add_abstract_numbering(numbering_elm, 1, heading_levels)
    _add_num_instance(numbering_elm, BODY_NUM_ID, 1)

    # --- Lists: own definition so counters never interleave with headings.
    # Each level shows only its own counter, indented one step per level.
    _add_numbering_definition(numbering_elm, 2, LIST_NUM_ID, [
        _make_level(0, "decimal", "%1.", num_pos=720, tab_pos=1440, wrap_pos=0),
        _make_level(1, "decimal", "%2.", num_pos=1440, tab_pos=2160, wrap_pos=0),
        _make_level(2, "decimal", "%3.", num_pos=2160, tab_pos=2880, wrap_pos=0),
    ])

    heading_num_ids = [BODY_NUM_ID]
    for i in range(extra_regions):
        num_id = _FIRST_REGION_NUM_ID + i
        _add_num_instance(numbering_elm, num_id, 1, restart=True,
                          level_count=HEADING_LEVELS)
        heading_num_ids.append(num_id)

    return heading_num_ids, LIST_NUM_ID


def apply_numbering(paragraph, num_id: int, ilvl: int):
    """Apply numbering to a paragraph via its XML numPr element."""
    pPr = paragraph._element.get_or_add_pPr()

    # Remove any existing numPr
    existing = pPr.find(qn("w:numPr"))
    if existing is not None:
        pPr.remove(existing)

    numPr = OxmlElement("w:numPr")

    ilvl_elm = OxmlElement("w:ilvl")
    ilvl_elm.set(qn("w:val"), str(ilvl))
    numPr.append(ilvl_elm)

    numId_elm = OxmlElement("w:numId")
    numId_elm.set(qn("w:val"), str(num_id))
    numPr.append(numId_elm)

    pPr.append(numPr)


# ---------------------------------------------------------------------------
# Content classification
# ---------------------------------------------------------------------------

HEADING_PATTERNS = re.compile(
    r"^("
    r"article\s|section\s|recitals?|definitions?|exhibit\s|schedule\s|"
    r"appendix\s|annex\s|preamble|background|introduction|terms and conditions"
    r")",
    re.IGNORECASE,
)

# A section number: "1.", "1.1", "2.3.4", "1.1.1.1". At least one dot is
# always required, so prose opening with a bare figure ("30 days from notice")
# never reads as a section. The unpunctuated form additionally needs a capital
# after it, which keeps "2.5 percent per annum" out when it appears as a short
# bold line. (Same convention as check_document._HEADING_NUM and
# extract_redlines.HEADING_NUM_RE.)
SECTION_NUMBER = re.compile(
    r"^(?:"
    r"(\d{1,3}(?:\.\d{1,3})*)[.)]\s+"        # "1. " / "1.1. " / "1) "
    r"|(\d{1,3}(?:\.\d{1,3})+)\s+(?=[A-Z])"  # "1.1 Foo" — needs the capital
    r")"
)

# Boilerplate that is short, bold, and often ALL CAPS but is never a section:
# cover legends, privilege banners, execution markers.
NEVER_NUMBER = re.compile(
    r"^("
    r"draft|confidential|privileged|subject to contract|for discussion|"
    r"attorney[- ]client|work product|not for distribution|execution (copy|version)"
    r")",
    re.IGNORECASE,
)

# Numbering this script cannot express natively but must not double up on:
# "ARTICLE I", "Section 1.1", "IV.", "C.". Broader than SECTION_NUMBER, and
# used only to detect that a document already numbers itself.
# The letter and roman branches require a trailing "." or ")" so that an
# ordinary heading ("RECITALS", "A MASTER SERVICES AGREEMENT") is not read as
# numbered on the strength of its first letter.
EXISTING_NUMBERING = re.compile(
    r"^(?:"
    r"(?:article|section|clause|part)\s+\S"   # ARTICLE I, Section 1.1
    r"|\d{1,3}(?:\.\d{1,3})*[.)\s]"           # 1.  1.1  1)
    r"|[ivxlc]+[.)]"                          # IV.  iv)
    r"|[a-z][.)]"                             # A.  a)
    r")",
    re.IGNORECASE,
)

# Headings that open a fresh numbering region — an exhibit's "1. Scope" is a
# deliberate restart, not a numbering mistake.
EXHIBIT_BOUNDARY = re.compile(
    r"^(exhibit|schedule|annex|appendix)\b",
    re.IGNORECASE,
)

# Body-text cross-references, used only to warn when renumbering may have
# invalidated them.
CROSS_REFERENCE = re.compile(
    r"\b(?:section|article|clause)\s+\d{1,3}(?:\.\d{1,3})*",
    re.IGNORECASE,
)

NUMBERED_PREFIX = re.compile(
    r"^(\d{1,3}\.(\d{1,3}\.)*\s+|\(\d{1,3}\)\s+|\([a-z]\)\s+|\([ivxlc]+\)\s+|[a-z]\.\s+|[ivxlc]+\.\s+)",
    re.IGNORECASE,
)

# Letter/roman-numeral markers ("a.", "iv.") need extra checks because they
# also match abbreviations ("E. coli") and case citations ("v. Smith").
LETTER_PREFIX = re.compile(r"^([a-z]|[ivxlc]+)\.\s+", re.IGNORECASE)

BULLET_CHARS = set("•·–—-*▪▸►")


def match_number_prefix(text: str):
    """Match a numbered-list prefix, or None.

    Digit forms ("1.", "2.3.", "(1)") always count. Letter/roman forms
    ("a.", "iv.") count only when the remainder does not start with a
    lowercase word (rules out abbreviations like "E. coli") and the marker
    is not "v"/"vs" before a capitalized word (rules out case citations
    like "v. Smith").
    """
    m = NUMBERED_PREFIX.match(text)
    if not m:
        return None
    letter = LETTER_PREFIX.match(text)
    if letter and letter.end() == m.end():
        rest = text[m.end():]
        if rest[:1].islower():
            return None
        if letter.group(1).lower() in ("v", "vs") and rest[:1].isupper():
            return None
    return m


def parse_section_number(text: str):
    """Split a leading section number off a heading.

    Returns (number, remainder) where number is a tuple of ints — "1.1" gives
    (1, 1) — or (None, text) when the heading carries no section number.

    Unlike match_number_prefix this accepts a final segment with no trailing
    dot, so "1.1 Defined Terms" is recognised. Without that, such headings
    were neither stripped nor skipped and came out double-numbered.
    """
    m = SECTION_NUMBER.match(text)
    if not m:
        return None, text
    digits = m.group(1) or m.group(2)
    return tuple(int(p) for p in digits.split(".")), text[m.end():].lstrip()


def format_section_number(number) -> str:
    """Render a number tuple the way Word will: (2, 1) -> '2.1.'"""
    return ".".join(str(n) for n in number) + "."


def _is_heading_style(style_name: str) -> bool:
    if not style_name:
        return False
    sn = style_name.lower()
    return sn.startswith("heading") or sn in ("title", "subtitle")


def _heading_level(style_name: str) -> int:
    if not style_name:
        return 1
    sn = style_name.lower()
    if sn == "title":
        return 0
    for n in (1, 2, 3):
        if str(n) in sn:
            return n
    return 1


def _style_bold(style) -> bool:
    """Resolve bold from a style, walking the base_style chain.

    Returns False when no style in the chain sets bold (unresolvable).
    """
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        font = getattr(style, "font", None)
        bold = font.bold if font is not None else None
        if bold is not None:
            return bold
        style = style.base_style
    return False


def _all_runs_bold(paragraph) -> bool:
    if not paragraph.runs:
        return False
    # run.bold is None when bold is inherited from the paragraph style, so
    # fall back to the style chain rather than treating None as not-bold.
    style_bold = _style_bold(paragraph.style)
    return all(
        r.bold if r.bold is not None else style_bold
        for r in paragraph.runs
        if r.text.strip()
    )


def _is_all_caps_text(text: str) -> bool:
    alpha = [c for c in text if c.isalpha()]
    return len(alpha) > 2 and all(c.isupper() for c in alpha)


class Classified(NamedTuple):
    """How a paragraph should be styled, and whether it is a numbered section.

    `kind` and `number` are deliberately independent. Being a heading (style
    it bold, give it heading spacing) and being a section (feed it to the
    numbering counter) are separate questions, and most numbering defects come
    from paragraphs where the first is true and the second is not — draft
    legends, privilege banners, signature labels.
    """

    kind: str            # 'title' | 'heading' | 'legend' | 'bullet' | 'number' | 'body'
    level: int           # 0 for title/body/legend, 1+ for headings and lists
    number: tuple | None # literal section number, e.g. (1, 1) for "1.1"
    source: str          # how a heading was detected: 'style' | 'heuristic' | 'keyword'
    centered: bool


def _is_centered(paragraph) -> bool:
    return paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER


def classify_paragraph(paragraph) -> Classified:
    """Classify a paragraph for styling and numbering."""
    text = paragraph.text.strip()
    centered = _is_centered(paragraph)

    if not text:
        return Classified("body", 0, None, "", centered)

    # Cover legends and privilege banners are short, bold and often ALL CAPS,
    # which is exactly the heading heuristic's profile. Catch them first.
    if len(text) < 200 and NEVER_NUMBER.match(text):
        return Classified("legend", 0, None, "", centered)

    style_name = paragraph.style.name if paragraph.style else ""

    # Existing heading style
    if _is_heading_style(style_name):
        level = _heading_level(style_name)
        if level == 0:
            return Classified("title", 0, None, "style", centered)
        number, _ = parse_section_number(text)
        return Classified("heading", level, number, "style", centered)

    # Short + bold or ALL CAPS → heading (even with number prefix)
    if len(text) < 120 and (_all_runs_bold(paragraph) or _is_all_caps_text(text)):
        if text[0] not in BULLET_CHARS:
            number, _ = parse_section_number(text)
            return Classified("heading", 1, number, "heuristic", centered)

    # Heading keyword patterns
    if len(text) < 120 and HEADING_PATTERNS.match(text):
        return Classified("heading", 1, None, "keyword", centered)

    # Existing list styles
    sn = style_name.lower()
    if "bullet" in sn or "list bullet" in sn:
        level = 2 if "2" in sn else 3 if "3" in sn else 1
        return Classified("bullet", level, None, "", centered)
    if "list number" in sn or "list continue" in sn:
        level = 2 if "2" in sn else 3 if "3" in sn else 1
        return Classified("number", level, None, "", centered)

    # Detect bullet by leading character
    if text and text[0] in BULLET_CHARS:
        return Classified("bullet", 1, None, "", centered)

    # Detect numbered list by prefix
    if match_number_prefix(text):
        return Classified("number", 1, None, "", centered)

    return Classified("body", 0, None, "", centered)


def strip_bullet_prefix(text: str) -> str:
    if text and text[0] in BULLET_CHARS:
        return text[1:].lstrip()
    return text


def strip_number_prefix(text: str) -> str:
    m = match_number_prefix(text)
    if m:
        return text[m.end():]
    return text


# ---------------------------------------------------------------------------
# Document building
# ---------------------------------------------------------------------------


def _smartquotes(text: str, prev_char: str = "") -> str:
    """Convert straight quotes to curly (smart) quotes.

    prev_char is the last character of the preceding run in the same
    paragraph (empty at paragraph start), so a quote at position 0 of a run
    is not blindly treated as opening when Word split the run mid-sentence.
    """
    result = []
    for i, ch in enumerate(text):
        before = text[i - 1] if i > 0 else prev_char
        if ch == '"':
            if not before or before in " \t\n\r([":
                result.append("\u201c")
            else:
                result.append("\u201d")
        elif ch == "'":
            if not before or before in " \t\n\r([":
                # Apostrophe before a digit ('90s) is an elision, not an
                # opening quote.
                if text[i + 1 : i + 2].isdigit():
                    result.append("\u2019")
                else:
                    result.append("\u2018")
            else:
                result.append("\u2019")
        else:
            result.append(ch)
    return "".join(result)


def _remove_paragraph_border(paragraph):
    """Remove all paragraph borders (line under/over/beside)."""
    pPr = paragraph._element.get_or_add_pPr()
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        pBdr.append(border)
    pPr.append(pBdr)


def _set_docdefaults_font(doc):
    """Pin docDefaults to the document font.

    Anything that does not resolve a font through a style falls back here —
    most visibly Word's auto-number glyphs. Theme attributes (asciiTheme and
    friends) take precedence over explicit ones, so they are removed rather
    than merely overwritten; leaving them is what renders numbers in the
    template's Cambria or Calibri beside Times New Roman text.
    """
    styles = doc.styles.element

    defaults = styles.find(qn("w:docDefaults"))
    if defaults is None:
        defaults = OxmlElement("w:docDefaults")
        styles.insert(0, defaults)

    rpr_default = defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        defaults.insert(0, rpr_default)

    rPr = rpr_default.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        rpr_default.append(rPr)

    fonts = rPr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rPr.insert(0, fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attr), DEFAULT_FONT)
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        if fonts.get(qn(attr)) is not None:
            del fonts.attrib[qn(attr)]

    sz = rPr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rPr.append(sz)
    sz.set(qn("w:val"), str(int(DEFAULT_SIZE.pt * 2)))


def _set_style_defaults(doc):
    """Set default font and remove title border on document styles."""
    _set_docdefaults_font(doc)
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        try:
            style = doc.styles[style_name]
            style.font.name = DEFAULT_FONT
            style.font.size = DEFAULT_SIZE
        except KeyError:
            pass
    # Remove bottom border from Title style
    try:
        title_style = doc.styles["Title"]
        pPr = title_style.element.get_or_add_pPr()
        existing_bdr = pPr.find(qn("w:pBdr"))
        if existing_bdr is not None:
            pPr.remove(existing_bdr)
    except (KeyError, AttributeError):
        pass


def source_runs(source_paragraph):
    """Runs in document order, including runs inside hyperlinks.

    paragraph.runs yields only direct w:r children, so hyperlink text
    (w:hyperlink/w:r) is invisible to it. Copying only .runs silently drops a
    linked citation or URL from the output while paragraph.text — used for
    classification and counts — still reports it as present.
    """
    return [
        Run(r, source_paragraph)
        for r in source_paragraph._p.xpath("./w:r | ./w:hyperlink/w:r")
    ]


def copy_runs(source_paragraph, dest_paragraph):
    """Copy runs preserving bold/italic/underline only."""
    prev_char = ""
    for src_run in source_runs(source_paragraph):
        text = src_run.text
        if not text:
            continue
        dest_run = dest_paragraph.add_run(_smartquotes(text, prev_char))
        prev_char = text[-1]
        dest_run.font.name = DEFAULT_FONT
        dest_run.font.size = DEFAULT_SIZE
        dest_run.bold = src_run.bold
        dest_run.italic = src_run.italic
        dest_run.underline = src_run.underline
        if src_run.font.strike:
            dest_run.font.strike = src_run.font.strike
        if src_run.font.subscript:
            dest_run.font.subscript = src_run.font.subscript
        if src_run.font.superscript:
            dest_run.font.superscript = src_run.font.superscript


def copy_runs_strip_prefix(source_paragraph, dest_paragraph, strip_fn):
    """Copy runs, stripping a prefix from the first non-empty run."""
    first = True
    prev_char = ""
    for src_run in source_runs(source_paragraph):
        run_text = src_run.text
        if first and run_text.strip():
            run_text = strip_fn(run_text.lstrip())
            first = False
        if not run_text:
            continue
        dest_run = dest_paragraph.add_run(_smartquotes(run_text, prev_char))
        prev_char = run_text[-1]
        dest_run.font.name = DEFAULT_FONT
        dest_run.font.size = DEFAULT_SIZE
        dest_run.bold = src_run.bold
        dest_run.italic = src_run.italic
        dest_run.underline = src_run.underline


def copy_table(source_table, dest_doc, warnings):
    """Copy a table with clean formatting."""
    rows = len(source_table.rows)
    cols = len(source_table.columns)
    dest_table = dest_doc.add_table(rows=rows, cols=cols)

    try:
        dest_table.style = dest_doc.styles["Table Grid"]
    except KeyError:
        pass

    # row.cells repeats the same underlying tc element for merged spans;
    # track seen tc elements so merged content is copied only once.
    seen_tcs = set()
    merged_warned = False
    nested_tables_dropped = 0

    for i, src_row in enumerate(source_table.rows):
        for j, src_cell in enumerate(src_row.cells):
            tc_id = id(src_cell._tc)
            if tc_id in seen_tcs:
                if not merged_warned:
                    warnings.append(
                        "Merged table cell detected; content copied once and "
                        "the merge structure was not preserved"
                    )
                    merged_warned = True
                continue
            seen_tcs.add(tc_id)

            dest_cell = dest_table.cell(i, j)
            if dest_cell.paragraphs:
                dest_cell.paragraphs[0].clear()

            for k, src_para in enumerate(src_cell.paragraphs):
                if k == 0:
                    dest_para = dest_cell.paragraphs[0]
                else:
                    dest_para = dest_cell.add_paragraph()

                try:
                    dest_para.style = dest_doc.styles["Normal"]
                except KeyError:
                    pass

                copy_runs(src_para, dest_para)

            # A table nested inside a cell (signature blocks, fee schedules)
            # is not reachable via cell.paragraphs and would vanish silently.
            if src_cell.tables:
                nested_tables_dropped += len(src_cell.tables)

    if nested_tables_dropped:
        warnings.append(
            f"{nested_tables_dropped} table(s) nested inside table cells were "
            "dropped; clean_format does not preserve tables within cells"
        )

    return dest_table


def iter_block_elements(container, warnings):
    """Yield w:p and w:tbl elements in document order.

    Content controls (w:sdt — cover pages, TOCs, etc.) are unwrapped so the
    paragraphs and tables inside w:sdtContent are processed like top-level
    blocks instead of being silently dropped. Handles nested controls.
    """
    for element in container:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag in ("p", "tbl"):
            yield element
        elif tag == "sdt":
            warnings.append(
                "Content control (w:sdt) unwrapped; its contents were "
                "processed as body content"
            )
            content = element.find(qn("w:sdtContent"))
            if content is not None:
                yield from iter_block_elements(content, warnings)


def _local_tag(element) -> str:
    return element.tag.split("}")[-1] if "}" in element.tag else element.tag


def _strip_section_number(text: str) -> str:
    number, remainder = parse_section_number(text)
    return remainder if number is not None else text


class Item(NamedTuple):
    """One body block, classified during the survey pass."""

    element: object
    tag: str
    paragraph: object | None
    text: str
    cls: Classified | None


def survey_blocks(blocks, source) -> tuple[list, str]:
    """Pass 1: classify every paragraph and pick the document's numbering mode.

    A document with two or more headings carrying literal decimal section
    numbers is presumed to number itself deliberately, and is mirrored.

    A document that numbers itself in a form this script cannot express
    natively — "ARTICLE I", "Section 1.1" — is left alone entirely
    ('preserve'). Inventing numbers for it would stack a second scheme on top
    of the author's and render "1. ARTICLE I".

    Anything else (a memo, a letter) gets numbering invented for it as before.
    """
    items = []
    numbered_headings = 0
    prenumbered_headings = 0

    for element in blocks:
        tag = _local_tag(element)
        if tag != "p":
            items.append(Item(element, tag, None, "", None))
            continue
        paragraph = Paragraph(element, source)
        cls = classify_paragraph(paragraph)
        text = paragraph.text.strip()
        if cls.kind == "heading":
            if cls.number is not None:
                numbered_headings += 1
            elif EXISTING_NUMBERING.match(text):
                prenumbered_headings += 1
        items.append(Item(element, tag, paragraph, text, cls))

    if numbered_headings >= 2:
        mode = "mirror"
    elif prenumbered_headings >= 2:
        mode = "preserve"
    else:
        mode = "auto"

    return items, mode


def _should_number(cls: Classified, mode: str) -> bool:
    """Whether a paragraph feeds the section counter.

    Centred paragraphs are never sections. In mirror mode the literal number
    is the signal; in auto mode only a real Heading N style counts, so the
    bold/ALL-CAPS heuristic can style a banner without numbering it.
    """
    if cls.kind != "heading" or cls.centered:
        return False
    if mode == "preserve":
        return False
    if mode == "mirror":
        return cls.number is not None
    return cls.source == "style"


def plan_numbering(items, mode: str):
    """Pass 2: assign (region, ilvl) to each numbered section.

    Counters are authoritative — the emitted number is always the counter, so
    duplicates, backwards jumps and gaps normalise to a sequential run. The
    literal number is compared only to report what changed. An exhibit or
    schedule heading opens a fresh region whose counters restart at 1.
    """
    counters = [0] * HEADING_LEVELS
    plans: dict[int, tuple[int, int]] = {}
    corrections = []
    region = 0
    region_count = 1
    pending_restart = False

    for idx, item in enumerate(items):
        cls = item.cls
        if cls is None:
            continue

        if cls.kind == "heading" and EXHIBIT_BOUNDARY.match(item.text):
            pending_restart = True
            continue

        if not _should_number(cls, mode):
            continue

        if pending_restart:
            pending_restart = False
            region = region_count
            region_count += 1
            counters = [0] * HEADING_LEVELS

        if mode == "mirror":
            ilvl = min(len(cls.number) - 1, HEADING_LEVELS - 1)
        else:
            ilvl = min(max(cls.level, 1) - 1, HEADING_LEVELS - 1)

        # A document opening at "1.1.1" would otherwise emit 0.0.1.
        for ancestor in range(ilvl):
            if counters[ancestor] == 0:
                counters[ancestor] = 1

        counters[ilvl] += 1
        for deeper in range(ilvl + 1, HEADING_LEVELS):
            counters[deeper] = 0

        emitted = tuple(counters[:ilvl + 1])
        if mode == "mirror" and cls.number is not None and cls.number != emitted:
            corrections.append({
                "heading": item.text,
                "source": format_section_number(cls.number),
                "emitted": format_section_number(emitted),
            })

        plans[idx] = (region, ilvl)

    return plans, region_count, corrections


def _cross_reference_warning(items, corrections):
    """Warn when renumbering may have invalidated body-text cross-references.

    Renumbering section 9.4 to 3.1 leaves "as set out in Section 9.4" in the
    prose pointing at nothing. The script does not rewrite prose, so this is
    the only signal the operator gets.
    """
    if not corrections:
        return None
    body = " ".join(
        item.text for item in items
        if item.cls is not None and item.cls.kind in ("body", "legend")
    )
    refs = CROSS_REFERENCE.findall(body)
    if not refs:
        return None
    return (
        f"{len(corrections)} heading(s) were renumbered and the document "
        f"contains {len(refs)} body-text cross-reference(s); verify each "
        "cross-reference still points at the intended section"
    )


def build_clean_document(source_path, template_path):
    """Build a clean-formatted document from source using template styles."""
    source = Document(str(source_path))

    if template_path and Path(template_path).exists():
        dest = Document(str(template_path))
    else:
        dest = Document()

    # Remove any default content from template
    for p in dest.paragraphs:
        p._element.getparent().remove(p._element)

    # Set margins
    for section in dest.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # Set default font on styles and remove title border
    _set_style_defaults(dest)

    results = {
        "paragraphs_processed": 0,
        "headings_detected": 0,
        "lists_converted": 0,
        "tables_processed": 0,
        "numbering_mode": "auto",
        "corrections": [],
        "warnings": [],
    }

    # Pass 1: classify everything, then decide how the document is numbered.
    # iter_block_elements is consumed once so its warnings are not duplicated.
    blocks = list(iter_block_elements(source.element.body, results["warnings"]))
    items, mode = survey_blocks(blocks, source)
    plans, region_count, corrections = plan_numbering(items, mode)

    results["numbering_mode"] = mode
    results["corrections"] = corrections

    stale = _cross_reference_warning(items, corrections)
    if stale:
        results["warnings"].append(stale)

    # Create numbering definitions (headings and lists are independent)
    heading_num_ids, list_num_id = create_numbering(dest, extra_regions=region_count - 1)

    title_seen = False

    # Pass 2: emit.
    for idx, item in enumerate(items):
        if item.tag == "p":
            src_para = item.paragraph
            text = item.text
            cls = item.cls
            results["paragraphs_processed"] += 1

            if not text:
                dest.add_paragraph("", style="Normal")
                continue

            plan = plans.get(idx)
            kind = cls.kind

            # Promote the first unnumbered ALL CAPS heading to title. A
            # heading that carries any numbering of its own is a section, not
            # a title, and neither is an exhibit divider.
            if (
                not title_seen
                and kind == "heading"
                and plan is None
                and _is_all_caps_text(text)
                and not EXHIBIT_BOUNDARY.match(text)
                and not EXISTING_NUMBERING.match(text)
            ):
                kind = "title"

            if kind == "title":
                title_seen = True
                results["headings_detected"] += 1
                dest_para = dest.add_paragraph(style="Title")
                copy_runs(src_para, dest_para)
                _remove_paragraph_border(dest_para)

            elif kind == "legend":
                # Cover legends and privilege banners keep their emphasis and
                # centring but never join the section numbering.
                dest_para = dest.add_paragraph(style="Normal")
                copy_runs(src_para, dest_para)
                if cls.centered:
                    dest_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            elif kind == "heading":
                results["headings_detected"] += 1
                # A numbered section's depth comes from its number, so
                # "1. DEFINITIONS" is Heading 1 even if it was Heading 2 in
                # the source.
                style_level = (plan[1] + 1) if plan is not None else max(cls.level, 1)
                style_name = f"Heading {min(style_level, 3)}"
                try:
                    dest_para = dest.add_paragraph(style=style_name)
                except KeyError:
                    dest_para = dest.add_paragraph(style="Normal")
                if plan is not None:
                    region, ilvl = plan
                    # Strip the literal number and let Word render it.
                    copy_runs_strip_prefix(src_para, dest_para, _strip_section_number)
                    apply_numbering(dest_para, heading_num_ids[region], ilvl)
                else:
                    copy_runs(src_para, dest_para)
                if cls.centered:
                    dest_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            elif kind == "bullet":
                results["lists_converted"] += 1
                dest_para = dest.add_paragraph(style="Normal")
                copy_runs_strip_prefix(src_para, dest_para, strip_bullet_prefix)
                # List items use their own numbering definition: L1=0, L2=1, L3=2
                ilvl = min(cls.level - 1, 2)
                apply_numbering(dest_para, list_num_id, ilvl)

            elif kind == "number":
                results["lists_converted"] += 1
                dest_para = dest.add_paragraph(style="Normal")
                copy_runs_strip_prefix(src_para, dest_para, strip_number_prefix)
                # List items use their own numbering definition: L1=0, L2=1, L3=2
                ilvl = min(cls.level - 1, 2)
                apply_numbering(dest_para, list_num_id, ilvl)

            else:
                dest_para = dest.add_paragraph(style="Normal")
                copy_runs(src_para, dest_para)

        elif item.tag == "tbl":
            from docx.table import Table as DocxTable

            src_table = DocxTable(item.element, source)
            copy_table(src_table, dest, results["warnings"])
            results["tables_processed"] += 1

    # Copy headers/footers
    try:
        for i, src_section in enumerate(source.sections):
            if i >= len(dest.sections):
                break
            dest_section = dest.sections[i]
            if src_section.header and src_section.header.paragraphs:
                header_text = "\n".join(
                    p.text for p in src_section.header.paragraphs if p.text.strip()
                )
                if header_text:
                    dest_section.header.is_linked_to_previous = False
                    for p in dest_section.header.paragraphs:
                        p.clear()
                    dest_section.header.paragraphs[0].text = header_text
            if src_section.footer and src_section.footer.paragraphs:
                footer_text = "\n".join(
                    p.text for p in src_section.footer.paragraphs if p.text.strip()
                )
                if footer_text:
                    dest_section.footer.is_linked_to_previous = False
                    for p in dest_section.footer.paragraphs:
                        p.clear()
                    dest_section.footer.paragraphs[0].text = footer_text
    except Exception as e:
        results["warnings"].append(f"Header/footer copy failed: {e}")

    return dest, results


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    args = sys.argv[1:]
    template_path = None

    if "--template" in args:
        idx = args.index("--template")
        template_path = args[idx + 1]
        args = args[:idx] + args[idx + 2:]

    if len(args) != 2:
        print(f"Usage: {sys.argv[0]} <input.docx> <output.docx> [--template <path>]")
        sys.exit(1)

    input_path = Path(args[0])
    output_path = Path(args[1])

    if not input_path.exists():
        print(f"Error: Input document not found: {input_path}")
        sys.exit(1)

    if not template_path:
        default_template = Path(__file__).parent / "legal-template.docx"
        if default_template.exists():
            template_path = str(default_template)

    doc, results = build_clean_document(input_path, template_path)
    doc.save(str(output_path))

    print(json.dumps(results, indent=2))

    total = results["paragraphs_processed"]
    headings = results["headings_detected"]
    lists = results["lists_converted"]
    tables = results["tables_processed"]
    warns = len(results["warnings"])
    print(
        f"\nSummary: {total} paragraphs, {headings} headings, "
        f"{lists} lists, {tables} tables, {warns} warnings",
        file=sys.stderr,
    )


if __name__ == "__main__":
    main()
