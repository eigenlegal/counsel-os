#!/usr/bin/env python3
"""Apply redline changes to a .docx file and add comments.

Usage:
    python3 apply_redlines.py [--track] <original.docx> <redlines.json> <output.docx>

With --track, every replacement is written as native Word tracked changes
(w:del + w:ins with author and timestamp) instead of a plain edit, so the
output IS the redline — no Word Compare pass needed. Only the changed core of
each replacement is marked: the common prefix and suffix of current/proposed
are trimmed at word boundaries, and deleted segments keep their original run
formatting. Without --track, behavior is unchanged (plain replacement).

The redlines.json file should contain an array of objects:
    [
      {
        "current": "exact text to find",
        "proposed": "replacement text",
        "comment": "rationale for the change (or null)",
        "author": "Author Name",
        "match": {
          "location": "body[12]",
          "before": "optional immediately preceding text",
          "after": "optional immediately following text"
        }
      }
    ]

If `current` appears more than once, the script refuses to apply that item
unless `match` selects exactly one occurrence. Supported match selectors:
`location`, `paragraph_index`, `occurrence`, `before`, `after`, and `context`.

Every item is resolved against the ORIGINAL document, not the partially
edited one: all targets are located before any replacement is made, so
`occurrence` numbers and other selectors always mean what they meant in the
document the redlines were drafted from. Consequently an item cannot target
text introduced by an earlier item, and two items whose targets overlap
result in the later item being skipped, never a silently misplaced edit.
"""

import copy
import itertools
import json
import sys
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.run import Run

from xml_safety import UnsafeXmlError, safe_fromstring

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
UNSUPPORTED_PARTS = (
    ("word/footnotes.xml", "footnote", "footnote"),
    ("word/endnotes.xml", "endnote", "endnote"),
    ("word/comments.xml", "comment", "comment"),
)


@dataclass
class TextMatch:
    location: str
    occurrence: int
    start: int
    end: int
    text: str
    before: str
    after: str
    replaceable: bool
    paragraph: Optional[Any] = None
    paragraph_index: Optional[int] = None


def get_runs(paragraph):
    """Collect a paragraph's runs in document order, including runs inside
    hyperlinks (w:hyperlink) and tracked insertions (w:ins).

    paragraph.runs only yields direct w:r children, which makes hyperlink
    and tracked-insertion text invisible to find/replace. Tracked deletions
    (w:del) are deliberately NOT descended into — deleted text must stay
    invisible.
    """
    return [
        Run(r, paragraph)
        for r in paragraph._p.xpath("./w:r | ./w:hyperlink/w:r | ./w:ins/w:r")
    ]


def get_paragraph_text(paragraph):
    """Get the full text of a paragraph by joining all runs."""
    return "".join(run.text for run in get_runs(paragraph))


def find_occurrence_starts(text, current):
    """Return non-overlapping start offsets for current in text."""
    if not current:
        return []

    starts = []
    start = 0
    while True:
        index = text.find(current, start)
        if index == -1:
            return starts
        starts.append(index)
        start = index + len(current)


def context_before(text, start_idx, size=160):
    return text[max(0, start_idx - size) : start_idx]


def context_after(text, end_idx, size=160):
    return text[end_idx : end_idx + size]


def truncate(text, length=80):
    return text[:length] + "..." if len(text) > length else text


def replace_in_paragraph(paragraph, current, proposed, start_idx=None):
    """Replace `current` text with `proposed` in a paragraph, preserving formatting.

    Joins run text to find the match, then reconstructs runs so the first
    matched run gets the replacement text and subsequent matched runs are
    cleared. Formatting of the first matched run is preserved.

    Returns True if a replacement was made, False otherwise.
    """
    full_text = get_paragraph_text(paragraph)
    if start_idx is None:
        start_idx = full_text.find(current)
    if start_idx == -1 or full_text[start_idx : start_idx + len(current)] != current:
        return False

    end_idx = start_idx + len(current)
    runs = get_runs(paragraph)

    # Map character positions to runs
    char_pos = 0
    run_ranges = []  # [(start_char, end_char, run_index)]
    for i, run in enumerate(runs):
        run_start = char_pos
        run_end = char_pos + len(run.text)
        run_ranges.append((run_start, run_end, i))
        char_pos = run_end

    # Find which runs are affected
    affected_runs = []
    for run_start, run_end, run_idx in run_ranges:
        if run_end > start_idx and run_start < end_idx:
            affected_runs.append((run_start, run_end, run_idx))

    if not affected_runs:
        return False

    first_run_start, first_run_end, first_run_idx = affected_runs[0]
    first_run = runs[first_run_idx]

    # Text before the match in the first affected run
    prefix = first_run.text[: start_idx - first_run_start]

    # Text after the match in the last affected run
    last_run_start, last_run_end, last_run_idx = affected_runs[-1]
    last_run = runs[last_run_idx]
    suffix = last_run.text[end_idx - last_run_start :]

    # Set the first run to prefix + proposed + suffix (if first == last)
    # or prefix + proposed (if multiple runs)
    if first_run_idx == last_run_idx:
        first_run.text = prefix + proposed + suffix
    else:
        first_run.text = prefix + proposed
        # Clear intermediate runs
        for _, _, run_idx in affected_runs[1:-1]:
            runs[run_idx].text = ""
        # Set last run to just the suffix
        last_run.text = suffix

    return True


# ---------------------------------------------------------------------------
# Native tracked changes (--track)
# ---------------------------------------------------------------------------


def split_replacement(current: str, proposed: str):
    """Trim the common prefix/suffix of a replacement at word boundaries.

    Returns (prefix_len, deleted_core, inserted_core). Marking only the
    changed core is what makes the redline readable: "30 days" -> "45 days"
    should strike "30" and insert "45", not the whole sentence.

    A boundary is widened only when it would split a word — a word character
    on the kept side AND on a core side ("30" -> "35" must never render as a
    struck "0"). Punctuation boundaries stand, so appending before a period
    is a pure insertion that keeps the period.
    """

    def word(s: str, i: int) -> bool:
        return 0 <= i < len(s) and s[i].isalnum()

    limit = min(len(current), len(proposed))

    prefix = 0
    while prefix < limit and current[prefix] == proposed[prefix]:
        prefix += 1
    while prefix > 0 and word(current, prefix - 1) and (
        word(current, prefix) or word(proposed, prefix)
    ):
        prefix -= 1

    suffix = 0
    max_suffix = limit - prefix
    while (
        suffix < max_suffix
        and current[len(current) - 1 - suffix] == proposed[len(proposed) - 1 - suffix]
    ):
        suffix += 1
    while suffix > 0 and word(current, len(current) - suffix) and (
        word(current, len(current) - suffix - 1)
        or word(proposed, len(proposed) - suffix - 1)
    ):
        suffix -= 1

    return prefix, current[prefix : len(current) - suffix], proposed[prefix : len(proposed) - suffix]


def make_revision_id_allocator(document):
    """Unique w:id values for new revision elements, above any existing."""
    max_id = 0
    for el in document.element.body.iter(qn("w:ins"), qn("w:del")):
        try:
            max_id = max(max_id, int(el.get(qn("w:id"), "0")))
        except ValueError:
            pass
    return itertools.count(max_id + 1)


def _revision_element(tag: str, author: str, when: str, id_alloc):
    el = OxmlElement(tag)
    el.set(qn("w:id"), str(next(id_alloc)))
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), when)
    return el


def _split_run(paragraph, run, offset: int):
    """Split a run at a text offset into (left, right), preserving rPr.

    The Run.text setter rebuilds content as w:t/w:tab/w:br, so tabs and
    breaks in either half round-trip correctly.
    """
    left_el = copy.deepcopy(run._r)
    run._r.addprevious(left_el)
    left = Run(left_el, paragraph)
    full = run.text
    left.text = full[:offset]
    run.text = full[offset:]
    return left, run


def _new_ins_run(paragraph, text: str, template_rpr):
    """A fresh w:r carrying `text`, formatted like the run it replaces."""
    r = OxmlElement("w:r")
    if template_rpr is not None:
        r.append(copy.deepcopy(template_rpr))
    Run(r, paragraph).text = text
    return r


def tracked_replace_in_paragraph(paragraph, current, proposed, start_idx,
                                 author, when, id_alloc):
    """Apply one replacement as native tracked changes (w:del + w:ins).

    Returns 'ok', 'not_found', or 'nested'. 'nested' means the changed core
    touches runs inside a hyperlink or an existing tracked insertion — OOXML
    forbids w:ins inside w:ins, and revision markup inside hyperlinks is not
    supported here, so the item is refused rather than emitted invalid.
    """
    full_text = get_paragraph_text(paragraph)
    if start_idx is None:
        start_idx = full_text.find(current)
    if start_idx == -1 or full_text[start_idx : start_idx + len(current)] != current:
        return "not_found"

    prefix_len, del_core, ins_core = split_replacement(current, proposed)
    if not del_core and not ins_core:
        return "ok"  # the edit is a no-op

    core_start = start_idx + prefix_len
    core_end = core_start + len(del_core)

    runs = get_runs(paragraph)
    char_pos = 0
    run_ranges = []
    for run in runs:
        run_ranges.append((char_pos, char_pos + len(run.text), run))
        char_pos += len(run.text)

    if del_core:
        affected = [
            (s, e, r) for s, e, r in run_ranges if e > core_start and s < core_end
        ]
    else:
        # Pure insertion: anchor on the run containing the insertion point,
        # preferring the run the point ends in over the one it starts.
        affected = [(s, e, r) for s, e, r in run_ranges if s <= core_start <= e]
        affected = affected[-1:]

    if not affected and not del_core and not run_ranges:
        # Insertion into an empty paragraph.
        ins_el = _revision_element("w:ins", author, when, id_alloc)
        ins_el.append(_new_ins_run(paragraph, ins_core, None))
        paragraph._p.append(ins_el)
        return "ok"
    if not affected:
        return "not_found"

    # Revision markup rearranges elements under the paragraph, which is only
    # coherent for runs that are direct children of it.
    for _, _, run in affected:
        if run._r.getparent() is not paragraph._p:
            return "nested"

    template_rpr = affected[0][2]._r.find(qn("w:rPr"))

    if not del_core:
        # Split the anchor run at the point and slot the w:ins between.
        run_start, _, run = affected[0]
        left, _right = _split_run(paragraph, run, core_start - run_start)
        ins_el = _revision_element("w:ins", author, when, id_alloc)
        ins_el.append(_new_ins_run(paragraph, ins_core, template_rpr))
        left._r.addnext(ins_el)
        return "ok"

    # Carve the deleted core out to whole runs.
    first_start, _, first_run = affected[0]
    if core_start > first_start:
        _split_run(paragraph, first_run, core_start - first_start)
    last_start, last_end, last_run = affected[-1]
    if core_end < last_end:
        # If first and last are the same run, its text now begins at
        # core_start after the split above.
        begins_at = core_start if last_run is first_run and core_start > first_start else last_start
        _split_run(paragraph, last_run, core_end - begins_at)

    # Re-derive the core's run elements after the splits.
    runs = get_runs(paragraph)
    char_pos = 0
    core_elements = []
    for run in runs:
        run_start, run_end = char_pos, char_pos + len(run.text)
        char_pos = run_end
        if run_start >= core_start and run_end <= core_end and run_end > run_start:
            core_elements.append(run._r)
    if not core_elements:
        return "not_found"

    del_el = _revision_element("w:del", author, when, id_alloc)
    core_elements[0].addprevious(del_el)
    for r in core_elements:
        del_el.append(r)  # moves the element
        for t in r.findall(qn("w:t")):
            t.tag = qn("w:delText")

    if ins_core:
        ins_el = _revision_element("w:ins", author, when, id_alloc)
        ins_el.append(_new_ins_run(paragraph, ins_core, template_rpr))
        del_el.addnext(ins_el)

    return "ok"


def collect_matches_from_paragraph(paragraph, current, location, occurrence, paragraph_index=None):
    full_text = get_paragraph_text(paragraph)
    matches = []

    for start in find_occurrence_starts(full_text, current):
        end = start + len(current)
        matches.append(
            TextMatch(
                location=location,
                occurrence=occurrence + len(matches),
                start=start,
                end=end,
                text=full_text,
                before=context_before(full_text, start),
                after=context_after(full_text, end),
                replaceable=True,
                paragraph=paragraph,
                paragraph_index=paragraph_index,
            )
        )

    return matches


def collect_replaceable_matches(document, current):
    matches = []

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        matches.extend(
            collect_matches_from_paragraph(
                paragraph,
                current,
                f"body[{paragraph_index}]",
                len(matches),
                paragraph_index=paragraph_index,
            )
        )

    for table_index, table in enumerate(document.tables):
        # row.cells repeats the same underlying tc element for merged spans;
        # skip already-seen tc elements so one occurrence counts once.
        seen_tcs = set()
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                tc_id = id(cell._tc)
                if tc_id in seen_tcs:
                    continue
                seen_tcs.add(tc_id)
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    matches.extend(
                        collect_matches_from_paragraph(
                            paragraph,
                            current,
                            (
                                f"table[{table_index}].row[{row_index}]"
                                f".cell[{cell_index}].p[{paragraph_index}]"
                            ),
                            len(matches),
                        )
                    )

    return matches


def paragraph_text_from_xml(paragraph):
    return "".join(node.text or "" for node in paragraph.iter(f"{W_NS}t"))


def collect_unsupported_matches(docx_path, current, occurrence):
    matches = []

    try:
        with zipfile.ZipFile(docx_path) as package:
            names = set(package.namelist())
            part_specs = list(UNSUPPORTED_PARTS)
            part_specs.extend(
                (name, "header", Path(name).stem)
                for name in names
                if name.startswith("word/header") and name.endswith(".xml")
            )
            part_specs.extend(
                (name, "footer", Path(name).stem)
                for name in names
                if name.startswith("word/footer") and name.endswith(".xml")
            )

            for part_name, part_kind, part_label in part_specs:
                if part_name not in names:
                    continue

                try:
                    root = safe_fromstring(package.read(part_name))
                except UnsafeXmlError as exc:
                    print(f"warning: skipping {part_name}: {exc}", file=sys.stderr)
                    continue
                paragraph_index = 0
                for paragraph in root.iter(f"{W_NS}p"):
                    full_text = paragraph_text_from_xml(paragraph)
                    for start in find_occurrence_starts(full_text, current):
                        end = start + len(current)
                        matches.append(
                            TextMatch(
                                location=f"{part_label}[{paragraph_index}]",
                                occurrence=occurrence + len(matches),
                                start=start,
                                end=end,
                                text=full_text,
                                before=context_before(full_text, start),
                                after=context_after(full_text, end),
                                replaceable=False,
                            )
                        )
                    paragraph_index += 1
    except (zipfile.BadZipFile, ET.ParseError):
        return matches

    return matches


def collect_matches(document, docx_path, current):
    matches = collect_replaceable_matches(document, current)
    matches.extend(collect_unsupported_matches(docx_path, current, len(matches)))
    return matches


def format_match(match):
    return {
        "location": match.location,
        "occurrence": match.occurrence,
        "start": match.start,
        "replaceable": match.replaceable,
        "before": truncate(match.before.strip(), 80),
        "after": truncate(match.after.strip(), 80),
    }


def select_match(matches, match_spec):
    selected = matches

    if not match_spec:
        if len(matches) == 1:
            return matches[0], None
        return None, f"Found {len(matches)} matches; add a match disambiguator"

    if not isinstance(match_spec, dict):
        return None, "match must be an object"

    if "occurrence" in match_spec:
        try:
            occurrence = int(match_spec["occurrence"])
        except (TypeError, ValueError):
            return None, "match.occurrence must be an integer"
        selected = [match for match in selected if match.occurrence == occurrence]

    if "location" in match_spec:
        selected = [
            match for match in selected if match.location == str(match_spec["location"])
        ]

    if "paragraph_index" in match_spec:
        try:
            paragraph_index = int(match_spec["paragraph_index"])
        except (TypeError, ValueError):
            return None, "match.paragraph_index must be an integer"
        selected = [
            match for match in selected if match.paragraph_index == paragraph_index
        ]

    if "before" in match_spec:
        before = str(match_spec["before"])
        selected = [match for match in selected if match.before.endswith(before)]

    if "after" in match_spec:
        after = str(match_spec["after"])
        selected = [match for match in selected if match.after.startswith(after)]

    if "context" in match_spec:
        context = str(match_spec["context"])
        selected = [match for match in selected if context in match.text]

    if len(selected) == 1:
        return selected[0], None
    if len(selected) == 0:
        return None, "match disambiguator selected no matches"
    return None, f"match disambiguator still selected {len(selected)} matches"


def text_in_tracked_deletions(document, current):
    """Return True if `current` appears in tracked-deletion text (w:delText)."""
    deleted = "".join(
        node.text or "" for node in document.element.body.iter(f"{W_NS}delText")
    )
    return current in deleted


def add_comment_to_paragraph(document, paragraph, comment_text, author):
    """Add a Word comment anchored to the full paragraph.

    Uses python-docx 1.2.0's native comment API with mark_comment_range
    to link the comment to the paragraph's runs. Anchors to direct runs
    when available, falling back to runs inside hyperlinks/tracked
    insertions. Returns True if the comment was added, False if the
    paragraph has no runs to anchor to.
    """
    runs = paragraph.runs or get_runs(paragraph)
    if not runs:
        return False

    comment = document.part.comments.add_comment(
        text=comment_text,
        author=author,
        initials="".join(word[0] for word in author.split() if word).upper(),
    )

    first_run = runs[0]
    last_run = runs[-1]
    first_run.mark_comment_range(last_run, comment.comment_id)
    return True


def main():
    args = [a for a in sys.argv[1:] if a != "--track"]
    track = "--track" in sys.argv[1:]
    if len(args) != 3:
        print(f"Usage: {sys.argv[0]} [--track] <original.docx> <redlines.json> <output.docx>")
        sys.exit(1)

    original_path = Path(args[0])
    redlines_path = Path(args[1])
    output_path = Path(args[2])

    if not original_path.exists():
        print(f"Error: Original document not found: {original_path}")
        sys.exit(1)

    if not redlines_path.exists():
        print(f"Error: Redlines JSON not found: {redlines_path}")
        sys.exit(1)

    with open(redlines_path) as f:
        redlines = json.load(f)

    doc = Document(str(original_path))

    results = {"applied": [], "skipped": [], "warnings": [], "tracked": track}

    revision_when = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    id_alloc = make_revision_id_allocator(doc) if track else None

    # Phase 1: resolve every item against the pristine document, before any
    # replacement mutates it. Occurrence numbers and offsets therefore always
    # refer to the original text the redlines were drafted from.
    resolved = []
    for i, item in enumerate(redlines):
        current = item["current"]
        match_spec = item.get("match")

        if not current:
            results["skipped"].append(
                {
                    "index": i,
                    "current": "",
                    "reason": "current text must not be empty",
                }
            )
            continue

        matches = collect_matches(doc, original_path, current)
        if not matches:
            if text_in_tracked_deletions(doc, current):
                results["warnings"].append(
                    {
                        "index": i,
                        "current": truncate(current),
                        "warning": (
                            "Text appears only inside tracked deletions (w:del); "
                            "deleted text is not editable"
                        ),
                    }
                )
            results["skipped"].append(
                {
                    "index": i,
                    "current": truncate(current),
                    "reason": "Text not found in document",
                }
            )
            continue

        selected_match, reason = select_match(matches, match_spec)
        if selected_match is None:
            results["skipped"].append(
                {
                    "index": i,
                    "current": truncate(current),
                    "reason": reason,
                    "matches": [format_match(match) for match in matches],
                }
            )
            continue

        if not selected_match.replaceable:
            results["skipped"].append(
                {
                    "index": i,
                    "current": truncate(current),
                    "reason": f"Selected match is in unsupported content: {selected_match.location}",
                    "matches": [format_match(selected_match)],
                }
            )
            continue

        resolved.append((i, item, selected_match))

    # Phase 2: apply back-to-front (descending start offset) so an applied
    # edit can never shift the offset of one still to come. Edits in
    # different paragraphs don't interact, so one global sort suffices;
    # the sort is stable, so of two items with the same target the earlier
    # one wins and the later is skipped by the pre-replace text check.
    for i, item, selected_match in sorted(
        resolved, key=lambda entry: -entry[2].start
    ):
        current = item["current"]
        proposed = item["proposed"]
        comment_text = item.get("comment")
        author = item.get("author", "Unknown")

        if track:
            status = tracked_replace_in_paragraph(
                selected_match.paragraph, current, proposed, selected_match.start,
                author, revision_when, id_alloc,
            )
            if status == "nested":
                results["skipped"].append(
                    {
                        "index": i,
                        "current": truncate(current),
                        "reason": (
                            "Changed text lies inside a hyperlink or an existing "
                            "tracked insertion; nested revision markup is not "
                            "supported — resolve the earlier revision first or "
                            "apply without --track"
                        ),
                    }
                )
                continue
            replaced = status == "ok"
        else:
            replaced = replace_in_paragraph(
                selected_match.paragraph, current, proposed, selected_match.start
            )
        if replaced:
            if comment_text:
                anchored = add_comment_to_paragraph(
                    doc, selected_match.paragraph, comment_text, author
                )
                if not anchored:
                    results["warnings"].append(
                        {
                            "index": i,
                            "current": truncate(current),
                            "warning": (
                                "Comment skipped: paragraph has no runs to "
                                "anchor the comment to"
                            ),
                        }
                    )
            results["applied"].append(
                {
                    "index": i,
                    "location": selected_match.location,
                    "occurrence": selected_match.occurrence,
                }
            )
            continue

        results["skipped"].append(
            {
                "index": i,
                "current": truncate(current),
                "reason": (
                    "Text at the resolved location changed before this edit "
                    "was applied (an earlier item's replacement overlaps it)"
                ),
            }
        )

    for entries in results.values():
        if isinstance(entries, list):
            entries.sort(key=lambda entry: entry["index"])

    doc.save(str(output_path))

    # Print results as JSON for the caller to parse
    print(json.dumps(results, indent=2))

    applied = len(results["applied"])
    skipped = len(results["skipped"])
    warnings = len(results["warnings"])
    print(f"\nSummary: {applied} applied, {skipped} skipped, {warnings} warnings", file=sys.stderr)

    if skipped > 0:
        sys.exit(2)  # Partial success


if __name__ == "__main__":
    main()
